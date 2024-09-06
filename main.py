from PyQt6 import uic
from PyQt6 import QtWidgets, QtGui, QtCore
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QPixmap, QPainter, QPageLayout, QFont, QIcon, QCursor
from PyQt6.QtWidgets import QMainWindow, QTableWidgetItem, QWidget, QHBoxLayout, QLabel, QMessageBox
import sys, io, random,os, webbrowser
from PyQt6.QtPrintSupport import QPrintDialog, QPrinter
import background_remover, fake_person_generator, fake_person_details_generator, url_shortener
from docx import Document
from docx.shared import Inches

class Ui(QMainWindow):
    def __init__(self):
        super(Ui, self).__init__()
        uic.loadUi('static/resources/interface.ui', self)
        self.setWindowTitle("BFU Tool")
        self.setWindowIcon(QIcon("static/resources/icon.png"))
        self.setFixedSize(611, 334)
        self.setWindowFlags(self.windowFlags() & ~QtCore.Qt.WindowType.WindowMaximizeButtonHint)
        self.label_3.setPixmap(QtGui.QPixmap("static/img/bgrm.jpg"))
        self.label_3.setGeometry(QtCore.QRect(20, 20, 201, 201))
        self.label_4.setPixmap(QtGui.QPixmap("static/img/bgrm.jpg"))
        self.label_4.setGeometry(QtCore.QRect(387, 20, 201, 201))
        self.label_14 = QtWidgets.QLabel(self.tab)
        self.label_14.setGeometry(QtCore.QRect(0, 0, 0, 0))
        self.label_14.setObjectName("label_14")
        self.label_5.setPixmap(QtGui.QPixmap("static/img/bgrm.jpg"))
        self.label_21.setPixmap(QtGui.QPixmap("static/img/img2.png"))
        self.label_22.setPixmap(QtGui.QPixmap("static/img/img1.png"))
        self.label_20.setPixmap(QtGui.QPixmap("static/img/img3.png"))
        self.label_16.setPixmap(QtGui.QPixmap("static/img/img7.png"))
        self.label_17.setPixmap(QtGui.QPixmap("static/img/img6.png"))
        self.label_18.setPixmap(QtGui.QPixmap("static/img/img5.png"))
        self.label_19.setPixmap(QtGui.QPixmap("static/img/img4.png"))
        self.show()
        self.pushButton.clicked.connect(self.open_file_dialog)
        self.pushButton_7.clicked.connect(self.remove_bg)
        self.pushButton_2.clicked.connect(self.save_image)
        self.pushButton_8.clicked.connect(self.generate_fake_image)
        self.pushButton_5.clicked.connect(self.save_generated_photo)
        self.pushButton_9.clicked.connect(self.generate_fake_details)
        self.pushButton_4.clicked.connect(self.save_line_edits_to_txt)
        self.pushButton_3.clicked.connect(self.save_details_to_docx)
        self.pushButton_6.clicked.connect(self.get_shortened_url)
        self.pushButton_10.clicked.connect(self.copy_to_clipboard)
        self.label_16.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.label_16.mousePressEvent = self.open_link_github
        self.label_17.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.label_17.mousePressEvent = self.open_link_site
        self.label_18.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.label_18.mousePressEvent = self.open_link_tg
        self.label_19.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.label_19.mousePressEvent = self.open_link_mail
    def open_link_github(self, event):
        webbrowser.open("https://github.com/dordea/")

    def open_link_site(self, event):
        webbrowser.open("https://dordea.github.io/")
    
    def open_link_tg(self, event):
        webbrowser.open("https://dordea.t.me/")
    
    def open_link_mail(self, event):
        webbrowser.open("mailto:pavlik@duck.com")

    def open_file_dialog(self):

        file_dialog = QtWidgets.QFileDialog(self)
        file_dialog.setFileMode(QtWidgets.QFileDialog.FileMode.ExistingFiles)
        file_dialog.setNameFilter("Images (*.png *.xpm *.jpg *.jpeg *.bmp)")
        file_dialog.setViewMode(QtWidgets.QFileDialog.ViewMode.List)

        if file_dialog.exec():
            file_paths = file_dialog.selectedFiles()
            if file_paths:
                self.load_image(file_paths[0])

    def load_image(self, image_path):
        self.label_14.setText(image_path)
        pixmap = QtGui.QPixmap(image_path)
        self.label_3.setPixmap(pixmap.scaled(self.label_3.size(), QtCore.Qt.AspectRatioMode.KeepAspectRatio))


    def remove_bg(self):
         pixmap = background_remover.remove_background(self.label_14.text())
         self.label_4.setGeometry(QtCore.QRect(387, 20, 201, 201))
         self.label_4.setAutoFillBackground(False)
         self.label_4.setFrameShape(QtWidgets.QFrame.Shape.StyledPanel)
         self.label_4.setFrameShadow(QtWidgets.QFrame.Shadow.Plain)
         self.label_4.setPixmap(QtGui.QPixmap(self.label_14.text()))
         self.label_4.setPixmap(QtGui.QPixmap(pixmap))

    def save_image(self):
    # Deschidem fereastra de dialog pentru a selecta calea și numele fișierului
        file_path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save Image", f"{self.label_14.text()[:-4]+'_bg-removed'+'.png'}", "PNG Files (*.png)")

    # Dacă utilizatorul alege un fișier
        if file_path:
        # Recuperăm imaginea din label_4
            pixmap = self.label_4.pixmap()

            if pixmap:
            # Salvăm imaginea la calea specificată
                pixmap.save(file_path)
                self.label_3.setPixmap(QtGui.QPixmap("static/img/bgrm.jpg"))
                self.label_4.setPixmap(QtGui.QPixmap("static/img/bgrm.jpg"))
                QMessageBox.information(self, "Succes", "Image Saved !").setWindowIcon(QtGui.QIcon(''))
            else:
                QMessageBox.warning(self, "Eroare", "No image to save!").setWindowIcon(QtGui.QIcon(''))

    def generate_fake_image(self):
        pixmap = fake_person_generator.generate_fake_person_image()
        self.label_5.setPixmap(QtGui.QPixmap(pixmap))
    
    def save_generated_photo(self):
    # Obține numele folderului din lineEdit
        folder_name = self.lineEdit.text()
        if not folder_name:
            QMessageBox.warning(self, "Error", "Folder name is empty!")
            return

    # Deschide dialogul QFileDialog pentru a selecta locația în care se va crea folderul
        save_path = QtWidgets.QFileDialog.getExistingDirectory(self, "Select Directory to Save Folder")

    # Dacă utilizatorul nu anulează dialogul
        if save_path:
            full_folder_path = os.path.join(save_path, folder_name)

        # Verifică dacă folderul există, dacă nu, îl creează
            if not os.path.exists(full_folder_path):
                os.makedirs(full_folder_path)

        # Deschide dialogul pentru a selecta calea și numele fișierului de salvat
            file_path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save Generated Photo",
                                                             os.path.join(full_folder_path, "Fake_person.png"),
                                                             "PNG Files (*.png)")

        # Dacă utilizatorul alege un fișier
            if file_path:
            # Recuperăm imaginea din label_5
                pixmap = self.label_5.pixmap()

                if pixmap:
                # Salvăm imaginea la calea specificată
                    pixmap.save(file_path)
                    self.label_5.setPixmap(QtGui.QPixmap("static/img/bgrm.jpg"))
                    QMessageBox.information(self, "Succes", "Generated Photo Saved!")
                else:
                    QMessageBox.warning(self, "Error", "No image to save!")

    def generate_fake_details(self):
        details = fake_person_details_generator.generate_fake_person_details()
        self.lineEdit.setText(details[0])
        self.lineEdit_2.setText(details[1])
        self.lineEdit_3.setText(details[2])
        self.lineEdit_4.setText(details[3])
        self.lineEdit_5.setText(details[4])
        self.lineEdit_6.setText(details[5])
        self.lineEdit_7.setText(details[6])
        self.lineEdit_8.setText(details[7])

    def save_line_edits_to_txt(self):
        line_edits_dict = {
        'Full Name': self.lineEdit.text(),
        'Email': self.lineEdit_2.text(),
        'Phone Number': self.lineEdit_3.text(),
        'Adress': self.lineEdit_4.text(),
        'Zip Code': self.lineEdit_5.text(),
        'City': self.lineEdit_6.text(),
        'State': self.lineEdit_7.text(),
        'Country': self.lineEdit_8.text(),
    }
        folder_name = self.lineEdit.text()  # Folosim textul din lineEdit pentru a crea numele folderului
        if not folder_name:
            return
        save_path = QtWidgets.QFileDialog.getExistingDirectory(None, "Select Directory to Save Folder")
    
        if save_path:  # Dacă utilizatorul nu anulează dialogul
            full_folder_path = os.path.join(save_path, folder_name)
        
            if not os.path.exists(full_folder_path):
                os.makedirs(full_folder_path)
        
            file_path = os.path.join(full_folder_path, "Fake_person.txt")
        
            with open(file_path, 'w') as file:
                for key, value in line_edits_dict.items():
                    file.write(f"{key}: {value}\n")

    def save_details_to_docx(self):
    # Obține numele folderului din lineEdit
        folder_name = self.lineEdit.text()
        if not folder_name:
            QMessageBox.warning(self, "Error", "Folder name is empty!")
            return

    # Deschide dialogul QFileDialog pentru a selecta locația unde se va crea folderul
        save_path = QtWidgets.QFileDialog.getExistingDirectory(self, "Select Directory to Save Folder")

    # Dacă utilizatorul nu anulează dialogul
        if save_path:
            full_folder_path = os.path.join(save_path, folder_name)

        # Verifică dacă folderul există, dacă nu, îl creează
            if not os.path.exists(full_folder_path):
                os.makedirs(full_folder_path)

        # Crează un document DOCX
            doc = Document()
            doc.add_heading('Fake Person Details', level=1)

        # Adaugă datele din QLineEdit-uri în document
            line_edits_dict = {
            'Full Name': self.lineEdit.text(),
            'Email': self.lineEdit_2.text(),
            'Phone Number': self.lineEdit_3.text(),
            'Address': self.lineEdit_4.text(),
            'Zip Code': self.lineEdit_5.text(),
            'City': self.lineEdit_6.text(),
            'State': self.lineEdit_7.text(),
            'Country': self.lineEdit_8.text(),
        }

            for key, value in line_edits_dict.items():
                doc.add_paragraph(f'{key}: {value}')

        # Adaugă imaginea din label_5
            pixmap = self.label_5.pixmap()
            if pixmap:
            # Salvăm imaginea temporar pentru a o adăuga în DOCX
                temp_image_path = os.path.join(full_folder_path, "temp_image.png")
                pixmap.save(temp_image_path)

                doc.add_paragraph('Profile Picture:')
                doc.add_picture(temp_image_path, width=Inches(2.0))  # Ajustează dimensiunea imaginii dacă e necesar

            # Ștergem imaginea temporară
                os.remove(temp_image_path)
            else:
                doc.add_paragraph('No profile picture available.')

        # Salvăm documentul DOCX
            docx_path = os.path.join(full_folder_path, "Fake_person.docx")
            doc.save(docx_path)

            QMessageBox.information(self, "Succes", "Details saved to DOCX file!")

    def get_shortened_url(self):
        self.lineEdit_10.setText(url_shortener.shorten_url(self.lineEdit_9.text()))
        self.pushButton_10.setText("Copy")
    
    def copy_to_clipboard(self):
        clipboard = QtWidgets.QApplication.clipboard()
        clipboard.setText(self.lineEdit_10.text())
        self.pushButton_10.setText("Copied!")

if __name__ == "__main__":
    app = QtWidgets.QApplication([])
    app.setWindowIcon(QIcon("static/resources/icon.png"))
    window = Ui()
    sys.exit(app.exec())